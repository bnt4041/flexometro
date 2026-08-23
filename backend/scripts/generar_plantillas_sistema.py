"""Genera las 3 plantillas Word de sistema (Fase 39).

Se ejecuta a mano, una vez, cuando el vocabulario de claves cambie — no en
producción. Los `.docx` resultantes se guardan en el repo, en
`app/modules/presupuestos/plantillas_docx_defecto/`, y
`plantilla_docx_service.asegurar_plantillas_sistema()` los sube a MinIO en
cada arranque de la API si no existen ya.

Uso: `python scripts/generar_plantillas_sistema.py`
"""

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

DESTINO = Path(__file__).parent.parent / "app/modules/presupuestos/plantillas_docx_defecto"


def _imagen_logo_placeholder() -> BytesIO:
    """PNG de relleno para el hueco del logo — el texto alternativo
    ("logo") es lo único que importa: `replace_pic` la localiza por ahí y
    sustituye sus bytes por el logo real de la organización al generar el
    documento (ver `plantilla_docx_service.generar_documento`)."""
    from PIL import Image, ImageDraw

    imagen = Image.new("RGB", (300, 100), color="#E5E5E5")
    dibujo = ImageDraw.Draw(imagen)
    dibujo.rectangle([0, 0, 299, 99], outline="#999999", width=2)
    dibujo.text((90, 40), "LOGO", fill="#666666")
    buffer = BytesIO()
    imagen.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _insertar_logo(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(_imagen_logo_placeholder(), width=900000)
    # `add_picture` mete el `<pic:cNvPr>` dentro de este mismo run: se busca
    # ahí, no en todo el documento, para no tocar el texto alternativo de
    # ninguna otra imagen que ya hubiera.
    for cNvPr in run._element.iter(qn("pic:cNvPr")):
        cNvPr.set("descr", "logo")

CLAVES_REFERENCIA = [
    ("presupuesto.codigo / .nombre / .descripcion", "Datos generales del presupuesto"),
    ("presupuesto.fecha / .emplazamiento / .tipo_obra / .validez_dias / .notas / .estado", ""),
    ("cliente.razon_social / .nif / .direccion / .codigo_postal / .ciudad / .provincia", "Vacío si el presupuesto no tiene cliente"),
    ("cliente.email / .telefono", ""),
    ("totales.pem / .gastos_generales / .beneficio_industrial", "Del coste al total, según el método de cálculo"),
    ("totales.pec_sin_iva / .porcentaje_iva / .iva / .total / .margen / .margen_pct", ""),
    ("capitulos (lista): .codigo / .resumen / .importe / .importe_venta", "Un elemento por capítulo"),
    ("partidas (lista): .capitulo_codigo / .capitulo_resumen / .codigo / .resumen / .texto", "Un elemento por partida"),
    ("partidas: .unidad / .medicion / .precio / .precio_venta / .importe / .importe_venta", ""),
    ("lineas_medicion (lista): .capitulo_codigo / .partida_codigo / .partida_resumen", "Un elemento por línea de medición"),
    ("lineas_medicion: .comentario / .uds / .longitud / .anchura / .altura / .parcial", ""),
    ("conceptos (lista): .codigo / .resumen / .unidad / .precio", "Un elemento por concepto del cuadro de precios"),
    ("componentes_descompuesto (lista): .concepto_codigo / .concepto_resumen", "Un elemento por línea de descomposición"),
    ("componentes_descompuesto: .hijo_codigo / .hijo_resumen / .hijo_unidad / .rendimiento / .factor / .precio", ""),
    ("organizacion.nombre / .cif / .direccion / .codigo_postal / .ciudad / .provincia", "Se editan en Ajustes -> Empresa"),
    ("organizacion.telefono / .email / .web", ""),
    ("organizacion.linkedin / .instagram / .facebook / .twitter", ""),
    ("banco.nombre / .entidad / .iban / .bic", "La cuenta marcada como predeterminada en Ajustes -> Bancos y cajas"),
    ("Logo: inserta cualquier imagen y ponle de texto alternativo la palabra logo", "Clic derecho sobre la imagen -> Editar texto alternativo -> logo"),
    ("Filtros de formato: valor|eur -> 1.234,56    valor|num(3) -> 1.234,560", "Se escriben con una barra vertical, sin llaves, dentro de una clave real"),
]


def _cabecera(doc: Document, titulo: str) -> None:
    _insertar_logo(doc)
    doc.add_heading("{{ organizacion.nombre }}", level=2)
    doc.add_paragraph(
        "{{ organizacion.direccion }}, {{ organizacion.codigo_postal }} {{ organizacion.ciudad }}"
        " ({{ organizacion.provincia }})    CIF: {{ organizacion.cif }}"
    )
    doc.add_paragraph("{{ organizacion.telefono }}    {{ organizacion.email }}")
    doc.add_paragraph()
    doc.add_heading(titulo, level=1)
    p = doc.add_paragraph()
    p.add_run("Presupuesto ").bold = False
    p.add_run("{{ presupuesto.codigo }}").bold = True
    p.add_run(" — {{ presupuesto.nombre }}")
    doc.add_paragraph("Fecha: {{ presupuesto.fecha }}    Emplazamiento: {{ presupuesto.emplazamiento }}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Cliente: ").bold = True
    p.add_run("{{ cliente.razon_social }} ({{ cliente.nif }})")
    doc.add_paragraph("{{ cliente.direccion }}, {{ cliente.codigo_postal }} {{ cliente.ciudad }} ({{ cliente.provincia }})")
    doc.add_paragraph()


def _tabla_bucle(
    doc: Document,
    encabezados: list[str],
    variable: str,
    campos: list[str],
    campos_eur: frozenset[str] = frozenset(),
) -> None:
    """Tabla con la fila de encabezado fija y una fila plantilla que
    `docxtpl` repite por cada elemento de `variable` (sintaxis `{%tr %}` de
    docxtpl). El `for`/`endfor` van cada uno en su propia fila de control
    —nunca junto a los datos en la misma fila—: docxtpl borra esas dos filas
    de control y las sustituye por el `{% %}` de Jinja desnudo, que pasa a
    envolver la fila de datos de en medio. Mezclar el `for`/`endfor` con los
    datos en la misma fila se probó y falla (docxtpl solo reconoce la última
    etiqueta `{%tr %}` de la fila, la del `endfor`, y descarta el resto)."""
    tabla = doc.add_table(rows=4, cols=len(encabezados))
    tabla.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla.rows[0].cells, encabezados, strict=True):
        celda.paragraphs[0].add_run(texto).bold = True

    tabla.rows[1].cells[0].text = f"{{%tr for x in {variable} %}}"
    fila = tabla.rows[2].cells
    for i, campo in enumerate(campos):
        filtro = "|eur" if campo in campos_eur else ""
        fila[i].text = f"{{{{ x.{campo}{filtro} }}}}"
    tabla.rows[3].cells[0].text = "{%tr endfor %}"


def _tabla_claves(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Claves disponibles", level=1)
    doc.add_paragraph(
        "Referencia para diseñar tu propia plantilla: descarga esta, bórrala o "
        "adáptala en Word manteniendo las claves entre llaves dobles que necesites, "
        "y súbela de nuevo en Ajustes -> Plantillas de presupuesto."
    )
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Light List Accent 1"
    tabla.rows[0].cells[0].paragraphs[0].add_run("Clave").bold = True
    tabla.rows[0].cells[1].paragraphs[0].add_run("Nota").bold = True
    for clave, nota in CLAVES_REFERENCIA:
        fila = tabla.add_row().cells
        fila[0].text = clave
        fila[1].text = nota


def plantilla_presupuesto() -> Document:
    doc = Document()
    _cabecera(doc, "Presupuesto")
    _tabla_bucle(
        doc,
        ["Capítulo", "Código", "Resumen", "Ud", "Medición", "Precio", "Importe"],
        "partidas",
        ["capitulo_codigo", "codigo", "resumen", "unidad", "medicion", "precio", "importe"],
        campos_eur=frozenset({"precio", "importe"}),
    )
    doc.add_paragraph()
    for etiqueta, clave in [
        ("PEM (coste)", "totales.pem|eur"),
        ("Gastos generales", "totales.gastos_generales|eur"),
        ("Beneficio industrial", "totales.beneficio_industrial|eur"),
        ("PEC sin IVA", "totales.pec_sin_iva|eur"),
        ("IVA ({{ totales.porcentaje_iva }} %)", "totales.iva|eur"),
    ]:
        doc.add_paragraph(f"{etiqueta}: {{{{ {clave} }}}} €")
    p = doc.add_paragraph()
    p.add_run(f"TOTAL: {{{{ totales.total|eur }}}} €").bold = True
    p.runs[0].font.size = Pt(14)

    # Solo se imprime si la empresa tiene una cuenta marcada como
    # predeterminada; si no, `banco` llega vacío y el bloque entero
    # desaparece. `{%p ... %}` (no `{% ... %}`) para que docxtpl borre
    # también los párrafos de control, igual que `{%tr %}` en las tablas.
    doc.add_paragraph()
    doc.add_paragraph("{%p if banco %}")
    doc.add_paragraph("Forma de pago: ingreso en {{ banco.entidad }}")
    doc.add_paragraph("IBAN: {{ banco.iban }}")
    doc.add_paragraph("{%p endif %}")

    _tabla_claves(doc)
    return doc


def plantilla_mediciones() -> Document:
    doc = Document()
    _cabecera(doc, "Estado de mediciones")
    _tabla_bucle(
        doc,
        ["Capítulo", "Partida", "Comentario", "Uds", "Long.", "Anch.", "Alt.", "Parcial"],
        "lineas_medicion",
        ["capitulo_codigo", "partida_codigo", "comentario", "uds", "longitud", "anchura", "altura", "parcial"],
    )
    _tabla_claves(doc)
    return doc


def plantilla_descompuestos() -> Document:
    doc = Document()
    _cabecera(doc, "Cuadro de precios descompuesto")
    doc.add_heading("Conceptos", level=2)
    _tabla_bucle(
        doc,
        ["Código", "Resumen", "Ud", "Precio"],
        "conceptos",
        ["codigo", "resumen", "unidad", "precio"],
        campos_eur=frozenset({"precio"}),
    )
    doc.add_paragraph()
    doc.add_heading("Descomposición", level=2)
    _tabla_bucle(
        doc,
        ["Concepto", "Componente", "Ud", "Rendimiento", "Factor", "Precio"],
        "componentes_descompuesto",
        ["concepto_codigo", "hijo_resumen", "hijo_unidad", "rendimiento", "factor", "precio"],
        campos_eur=frozenset({"precio"}),
    )
    _tabla_claves(doc)
    return doc


if __name__ == "__main__":
    DESTINO.mkdir(parents=True, exist_ok=True)
    plantilla_presupuesto().save(DESTINO / "presupuesto.docx")
    plantilla_mediciones().save(DESTINO / "mediciones.docx")
    plantilla_descompuestos().save(DESTINO / "descompuestos.docx")
    print(f"Generadas en {DESTINO}")
